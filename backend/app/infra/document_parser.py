"""
Parser de documentos - Extração de texto de arquivos

Este módulo fornece funcionalidades para extrair texto de documentos
em diversos formatos, incluindo PDFs, Word, PowerPoint, texto e imagens.

Funcionalidades:
- Extração de texto de PDFs (digitais e escaneados com OCR)
- Extração de texto de documentos Word (.docx, .doc)
- Extração de texto de apresentações PowerPoint (.pptx, .ppt)
- Extração de texto de arquivos de texto (.txt, .md)
- OCR para imagens e PDFs escaneados (.png, .jpg, .tiff)
- Metadados ricos (páginas, elementos, estrutura)
- Fallback para métodos simples se Unstructured não disponível

Arquitetura:
- Usa Unstructured.io como método preferido (extração inteligente)
- Fallback para PyPDF2/python-docx se Unstructured não disponível
- Suporte a OCR usando Tesseract (opcional)
- Validação de tamanho de arquivo (máximo configurável)

Configuração (parser_config.py):
- USE_OCR: Habilita/desabilita OCR (default: False)
- STRATEGY: Estratégia de parsing (auto, hi_res, fast)
- OCR_LANGUAGES: Idiomas para OCR (ex: "por+eng")
- MAX_FILE_SIZE_MB: Tamanho máximo de arquivo (default: 10MB)
- USE_UNSTRUCTURED: Usa Unstructured.io ou fallback (default: True)

Formatos suportados:
- PDF: application/pdf
- Word: application/vnd.openxmlformats-officedocument.wordprocessingml.document, application/msword
- PowerPoint: application/vnd.openxmlformats-officedocument.presentationml.presentation, application/vnd.ms-powerpoint
- Texto: text/plain, text/markdown
- Imagens: image/png, image/jpeg, image/tiff (com OCR)
"""

from typing import Optional, Dict, Any
import io
import tempfile
import os
from pathlib import Path

from backend.app.logging_config import get_logger

# Importa configurações
try:
    from backend.app.infra.parser_config import (
        USE_OCR,
        STRATEGY,
        OCR_LANGUAGES,
        MAX_FILE_SIZE_MB,
        USE_UNSTRUCTURED
    )
except ImportError:
    # Valores padrão se arquivo de config não existir
    USE_OCR = False  # PADRÃO: SEM OCR (não precisa Tesseract)
    STRATEGY = "fast"
    OCR_LANGUAGES = None
    MAX_FILE_SIZE_MB = 10
    USE_UNSTRUCTURED = True

logger = get_logger(__name__)

# Importações opcionais (instalar apenas quando necessário)
try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    logger.warning(
        "⚠️ unstructured não instalado. "
        "Instale com: pip install unstructured[all-docs]"
    )
    UNSTRUCTURED_AVAILABLE = False


class DocumentParser:
    """
    Parser de documentos com suporte a múltiplos formatos.
    
    Esta classe fornece métodos para extrair texto de documentos
    em diversos formatos, usando Unstructured.io como método preferido
    e fallback para métodos simples se necessário.
    
    Métodos:
    - is_supported(): Verifica se um tipo MIME é suportado
    - get_extension(): Retorna extensão para um tipo MIME
    - parse_file(): Extrai texto de um arquivo
    
    Estratégias:
    - Unstructured.io (preferido): Extração inteligente com estrutura
    - Fallback simples: PyPDF2, python-docx para formatos básicos
    
    Configuração:
    - USE_OCR: Habilita OCR para PDFs escaneados e imagens
    - STRATEGY: Estratégia de parsing (auto, hi_res, fast)
    - MAX_FILE_SIZE: Tamanho máximo de arquivo em bytes
    
    Attributes:
        SUPPORTED_TYPES: Dict com tipos MIME suportados
        MAX_FILE_SIZE: Tamanho máximo de arquivo em bytes
        use_unstructured: Se True, usa Unstructured.io (preferido)
    """

    SUPPORTED_TYPES = {
        # PDFs
        "application/pdf": [".pdf"],

        # Word
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
        "application/msword": [".doc"],

        # PowerPoint
        "application/vnd.openxmlformats-officedocument.presentationml.presentation": [".pptx"],
        "application/vnd.ms-powerpoint": [".ppt"],

        # Texto
        "text/plain": [".txt"],
        "text/markdown": [".md"],

        # Imagens (com OCR)
        "image/png": [".png"],
        "image/jpeg": [".jpg", ".jpeg"],
        "image/tiff": [".tiff", ".tif"],
    }

    MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024  # Converte MB para bytes

    def __init__(self, use_unstructured: Optional[bool] = None):
        """
        Inicializa o parser.

        Args:
            use_unstructured: Se True, usa Unstructured.io (recomendado)
                            Se False, usa fallback simples
                            Se None, usa valor do parser_config.py
        """
        if use_unstructured is None:
            use_unstructured = USE_UNSTRUCTURED

        self.use_unstructured = use_unstructured and UNSTRUCTURED_AVAILABLE

        if use_unstructured and not UNSTRUCTURED_AVAILABLE:
            logger.warning(
                "Unstructured solicitado mas não disponível. "
                "Usando fallback simples."
            )

        # Log da configuração atual
        logger.info(
            f"Parser configurado: "
            f"Unstructured={'ON' if self.use_unstructured else 'OFF'}, "
            f"OCR={'ON' if USE_OCR else 'OFF'}, "
            f"Strategy={STRATEGY}, "
            f"Max={MAX_FILE_SIZE_MB}MB"
        )

    def is_supported(self, mime_type: str) -> bool:
        """Verifica se o tipo MIME é suportado."""
        return mime_type in self.SUPPORTED_TYPES

    def get_extension(self, mime_type: str) -> Optional[str]:
        """Retorna a extensão típica para um tipo MIME."""
        extensions = self.SUPPORTED_TYPES.get(mime_type)
        return extensions[0] if extensions else None

    def parse_file(
        self,
        file_data: bytes,
        filename: str,
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Parseia um arquivo e extrai texto e metadados.

        Args:
            file_data: Bytes do arquivo
            filename: Nome original do arquivo
            mime_type: Tipo MIME do arquivo

        Returns:
            Dict com:
            - text: Texto extraído
            - metadata: Metadados do documento
            - elements: Elementos estruturados (se Unstructured usado)

        Raises:
            ValueError: Se arquivo não suportado ou muito grande
        """
        # Validações
        if not self.is_supported(mime_type):
            raise ValueError(
                f"Tipo de arquivo não suportado: {mime_type}. "
                f"Suportados: {', '.join(self.SUPPORTED_TYPES.keys())}"
            )

        file_size = len(file_data)
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(
                f"Arquivo muito grande: {file_size / 1024 / 1024:.2f} MB. "
                f"Máximo: {self.MAX_FILE_SIZE / 1024 / 1024:.2f} MB"
            )

        logger.info(
            f"Parseando arquivo: {filename} "
            f"({mime_type}, {file_size / 1024:.2f} KB)"
        )

        # Parseia com Unstructured ou fallback
        if self.use_unstructured:
            return self._parse_with_unstructured(file_data, filename, mime_type)
        else:
            return self._parse_simple(file_data, filename, mime_type)

    def _parse_with_unstructured(
        self,
        file_data: bytes,
        filename: str,
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Parseia usando Unstructured.io (método preferido).

        Vantagens:
        - Extração inteligente de estrutura (títulos, parágrafos, tabelas)
        - OCR para PDFs escaneados e imagens
        - Suporte a layouts complexos
        - Metadados ricos
        """
        try:
            # Unstructured precisa de arquivo físico temporário
            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=Path(filename).suffix
            ) as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name

            try:
                # Particiona o documento
                # CONFIGURAÇÃO IMPORTANTE: ocr_languages só é usado se USE_OCR = True
                partition_kwargs = {
                    "filename": temp_path,
                    "strategy": STRATEGY,  # Vem do parser_config.py
                    "include_metadata": True,
                    "include_page_breaks": True,
                }

                # Adiciona OCR apenas se configurado
                if USE_OCR and OCR_LANGUAGES:
                    partition_kwargs["ocr_languages"] = OCR_LANGUAGES
                    logger.info(f"🔍 OCR ativado com idiomas: {OCR_LANGUAGES}")
                else:
                    logger.info("⚡ OCR desativado - processamento rápido")

                elements = partition(**partition_kwargs)

                # Extrai texto de todos os elementos
                text_parts = []
                metadata_list = []

                for element in elements:
                    text_parts.append(str(element))

                    # Coleta metadados interessantes
                    if hasattr(element, 'metadata') and element.metadata:
                        meta = {
                            "type": element.category if hasattr(element, 'category') else None,
                            "page": element.metadata.page_number if hasattr(element.metadata, 'page_number') else None,
                        }
                        metadata_list.append(meta)

                text = "\n\n".join(text_parts)

                logger.info(
                    f"✅ Unstructured: {len(elements)} elementos extraídos, "
                    f"{len(text)} caracteres"
                )

                return {
                    "text": text,
                    "metadata": {
                        "parser": "unstructured",
                        "elements_count": len(elements),
                        "char_count": len(text),
                        "elements": metadata_list[:10],  # Primeiros 10
                    },
                    "elements": elements  # Elementos completos para análise avançada
                }

            finally:
                # Remove arquivo temporário
                try:
                    os.unlink(temp_path)
                except:
                    pass

        except Exception as e:
            logger.error(f"❌ Erro no Unstructured: {e}")
            logger.info("Tentando fallback simples...")
            return self._parse_simple(file_data, filename, mime_type)

    def _parse_simple(
        self,
        file_data: bytes,
        filename: str,
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Fallback simples sem Unstructured.

        Limitações:
        - Sem OCR para PDFs escaneados
        - Extração básica de texto
        - Perde formatação e estrutura
        """
        try:
            # Texto puro
            if mime_type in ["text/plain", "text/markdown"]:
                text = file_data.decode("utf-8", errors="ignore")
                return {
                    "text": text,
                    "metadata": {"parser": "simple_text"},
                    "elements": []
                }

            # PDF com PyPDF2 (básico)
            elif mime_type == "application/pdf":
                try:
                    import PyPDF2
                    pdf_file = io.BytesIO(file_data)
                    reader = PyPDF2.PdfReader(pdf_file)

                    text_parts = []
                    for page in reader.pages:
                        text_parts.append(page.extract_text())

                    text = "\n\n".join(text_parts)

                    return {
                        "text": text,
                        "metadata": {
                            "parser": "pypdf2",
                            "pages": len(reader.pages)
                        },
                        "elements": []
                    }
                except ImportError:
                    raise ValueError(
                        "PyPDF2 não instalado. "
                        "Instale com: pip install PyPDF2"
                    )

            # DOCX com python-docx (básico)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                try:
                    import docx
                    doc_file = io.BytesIO(file_data)
                    doc = docx.Document(doc_file)

                    text_parts = [para.text for para in doc.paragraphs]
                    text = "\n\n".join(text_parts)

                    return {
                        "text": text,
                        "metadata": {
                            "parser": "python-docx",
                            "paragraphs": len(doc.paragraphs)
                        },
                        "elements": []
                    }
                except ImportError:
                    raise ValueError(
                        "python-docx não instalado. "
                        "Instale com: pip install python-docx"
                    )

            else:
                raise ValueError(
                    f"Fallback simples não suporta {mime_type}. "
                    "Instale unstructured: pip install unstructured[all-docs]"
                )

        except Exception as e:
            logger.error(f"❌ Erro no fallback simples: {e}")
            raise


# Instância global
document_parser = DocumentParser(use_unstructured=True)
